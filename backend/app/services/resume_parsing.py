"""Extract plain text from PDF / DOCX / TXT resume uploads.

Returned text is the corpus the TF-IDF matcher (services/matching.py) consumes via
`Resume.content_text`. Keep the extraction lossy but consistent — strip excess whitespace.
"""
from __future__ import annotations

import io
import re

from docx import Document
from pypdf import PdfReader

PDF_MIMES = {"application/pdf"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/msword",  # accept legacy .doc filename; python-docx can't open true .doc, see below
}
TXT_MIMES = {"text/plain"}

ALLOWED_EXTS = {".pdf", ".docx", ".txt"}
MAX_FILE_BYTES = 5 * 1024 * 1024  # 5 MB


class ResumeParseError(ValueError):
    """Raised when the uploaded file cannot be parsed into text."""


_whitespace_re = re.compile(r"[ \t]+")
_blank_line_re = re.compile(r"\n{3,}")


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _whitespace_re.sub(" ", text)
    text = _blank_line_re.sub("\n\n", text)
    return text.strip()


def extract_text_from_pdf(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
    except Exception as e:  # pypdf raises various subclasses
        raise ResumeParseError(f"Could not read PDF: {e}") from e
    parts: list[str] = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:
            # Skip unreadable pages rather than failing the whole upload.
            continue
    return _normalize("\n\n".join(p for p in parts if p))


def extract_text_from_docx(content: bytes) -> str:
    try:
        doc = Document(io.BytesIO(content))
    except Exception as e:
        raise ResumeParseError(
            "Could not read DOCX. (Note: legacy .doc files aren't supported — save as .docx or PDF.)"
        ) from e
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    table_text: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                table_text.append(row_text)
    combined = "\n".join(paragraphs + table_text)
    return _normalize(combined)


def extract_text_from_txt(content: bytes) -> str:
    try:
        return _normalize(content.decode("utf-8"))
    except UnicodeDecodeError:
        return _normalize(content.decode("latin-1", errors="replace"))


def extract_text(content: bytes, filename: str, mime_type: str | None) -> str:
    """Dispatch by mime type, fall back to extension. Returns normalized text."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    mime = (mime_type or "").lower()

    if mime in PDF_MIMES or ext == "pdf":
        return extract_text_from_pdf(content)
    if mime in DOCX_MIMES or ext == "docx":
        return extract_text_from_docx(content)
    if mime in TXT_MIMES or ext == "txt":
        return extract_text_from_txt(content)
    raise ResumeParseError(
        f"Unsupported file type '{mime or ext}'. Upload a PDF, DOCX, or TXT file."
    )
